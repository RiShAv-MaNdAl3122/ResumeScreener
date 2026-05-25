import io
import re
import PyPDF2
from docx import Document

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Extracts text from a PDF file provided as bytes.
    """
    pdf_file = io.BytesIO(file_bytes)
    reader = PyPDF2.PdfReader(pdf_file)
    text = ""
    for page in reader.pages:
        extracted_text = page.extract_text()
        if extracted_text:
            text += extracted_text + "\n"
    return text

def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extracts text from a DOCX file provided as bytes.
    """
    docx_file = io.BytesIO(file_bytes)
    document = Document(docx_file)
    
    text_chunks = []
    
    # 1. Extract regular paragraphs
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            text_chunks.append(paragraph.text.strip())
            
    # 2. Extract text from tables (many resumes use tables for layouts)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip():
                        text_chunks.append(paragraph.text.strip())
                        
    return "\n".join(text_chunks)



